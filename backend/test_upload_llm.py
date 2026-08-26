import asyncio
import httpx
import os
import docx

async def run_test():
    async with httpx.AsyncClient(base_url="http://localhost:8000", timeout=120.0) as client:
        # 1. Login
        login_data = {
            "username": "testuser@forr.com",
            "password": "securepassword123"
        }
        res = await client.post("/api/v1/auth/login", data=login_data)
        assert res.status_code == 200, f"Login failed: {res.text}"
        access_token = res.json()["access_token"]
        headers = {"Authorization": f"Bearer {access_token}"}

        # 2. Get business ID
        res = await client.get("/api/v1/businesses", headers=headers)
        businesses = res.json()
        business_id = businesses[0]["id"]

        # 3. Create dummy DOCX with unstructured text
        docx_path = "test_unstructured.docx"
        doc = docx.Document()
        doc.add_paragraph("Welcome to our store!")
        doc.add_paragraph("We are selling a fantastic Wireless Mouse today for only 25 bucks. It's black and very fast.")
        doc.add_paragraph("We also have a Mechanical Keyboard available. The price is $75. We have about 10 of these in stock. Excellent for typing.")
        doc.save(docx_path)

        # 4. Upload and Parse (this hits the Groq LLM)
        print("Uploading to LLM parser... this may take a few seconds.")
        with open(docx_path, "rb") as f:
            files = {"file": ("test_unstructured.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            res = await client.post(
                f"/api/v1/businesses/{business_id}/products/upload/parse",
                headers=headers,
                files=files
            )
        assert res.status_code == 200, f"Parse failed: {res.text}"
        parse_data = res.json()
        
        # Verify LLM extracted exactly the headers we want
        assert "name" in parse_data["headers"]
        assert len(parse_data["data"]) == 2
        
        products = parse_data["data"]
        print("LLM successfully extracted products:")
        for p in products:
            print(f"- {p.get('name')} | ${p.get('price')} | Qty: {p.get('quantity')}")
            assert p.get("name")
            assert p.get("price") > 0

        # 5. Send Bulk Request (Already mapped by LLM!)
        mapped_payload = []
        for p in products:
            mapped_payload.append({
                "name": p.get("name"),
                "price": float(p.get("price", 0)),
                "description": p.get("description", ""),
                "quantity": int(p.get("quantity", 0)),
                "category": p.get("category", "Electronics"),
                "status": "draft"
            })
        
        res = await client.post(
            f"/api/v1/businesses/{business_id}/products/bulk",
            headers=headers,
            json=mapped_payload
        )
        assert res.status_code == 201, f"Bulk insert failed: {res.text}"
        print("DOCX LLM Bulk insert successful!")

        # Cleanup
        os.remove(docx_path)

if __name__ == "__main__":
    import sys
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    asyncio.run(run_test())
